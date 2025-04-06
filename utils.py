# Utility functions for PDF processing, database interactions, etc. 

import os
import sqlite3
import json
import io # For creating in-memory files
import logging # Added for logging
from datetime import datetime
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from typing import List, Dict, Optional
from langchain.docstore.document import Document # Use this specific import path

# --- Export Libraries ---
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.units import inch
from docx import Document as DocxDocument
from docx.shared import Inches

# --- Configuration ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

DB_PATH = os.path.join("db", "chat_history.sqlite")
UPLOAD_FOLDER = os.path.join("data", "uploaded_pdfs")
DEFAULT_CHUNK_SIZE = 1000
DEFAULT_CHUNK_OVERLAP = 150

# --- Database Functions ---

def init_db():
    """Initializes the SQLite database and creates/updates the exams table."""
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    try:
        with sqlite3.connect(DB_PATH) as conn:
            cursor = conn.cursor()

            # Create table if not exists
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS exams (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                    session_id TEXT,
                    pdf_name TEXT,
                    topic TEXT NOT NULL,
                    num_questions INTEGER NOT NULL,
                    questions_json TEXT,
                    answers_json TEXT
                )
            """)

            # Add session_id column if missing (backward compatibility)
            cursor.execute("PRAGMA table_info(exams)")
            columns = [col[1] for col in cursor.fetchall()]
            if 'session_id' not in columns:
                cursor.execute("ALTER TABLE exams ADD COLUMN session_id TEXT")
                logging.info("Added 'session_id' column to existing exams table.")

            conn.commit()
            logging.info("Database initialized successfully.")
    except sqlite3.Error as e:
        logging.error(f"Database initialization or schema alteration error: {e}", exc_info=True)

def save_exam(session_id: str, pdf_name: Optional[str], topic: str, num_questions: int, questions_json: str, answers_json: str):
    """Saves the generated exam details to the database."""
    try:
        with sqlite3.connect(DB_PATH) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO exams (session_id, pdf_name, topic, num_questions, questions_json, answers_json)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (session_id, pdf_name, topic, num_questions, questions_json, answers_json))
            conn.commit()
            logging.info(f"Exam saved successfully for topic: {topic}")
    except sqlite3.Error as e:
        logging.error(f"Error saving exam to database: {e}", exc_info=True)

def get_exam_history(limit: int = 10) -> List[dict]:
    """Retrieves the most recent exam history from the database."""
    history = []
    try:
        with sqlite3.connect(DB_PATH) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            cursor.execute("SELECT id, timestamp, session_id, pdf_name, topic, num_questions FROM exams ORDER BY timestamp DESC LIMIT ?", (limit,))
            history = [dict(row) for row in cursor.fetchall()]
    except sqlite3.Error as e:
        logging.error(f"Error retrieving exam history: {e}", exc_info=True)
    return history

def get_exam_by_id(exam_id: int) -> Optional[dict]:
    """Retrieves a specific exam record by its ID."""
    exam_data = None
    try:
        with sqlite3.connect(DB_PATH) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM exams WHERE id = ?", (exam_id,))
            row = cursor.fetchone()
            if row:
                exam_data = dict(row)
    except sqlite3.Error as e:
        logging.error(f"Error retrieving exam with ID {exam_id}: {e}", exc_info=True)
    return exam_data

def delete_exam(exam_id: int) -> bool:
    """Deletes an exam record from the database by its ID."""
    try:
        with sqlite3.connect(DB_PATH) as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM exams WHERE id = ?", (exam_id,))
            conn.commit()
            logging.info(f"Deleted exam record with ID: {exam_id}")
            return True
    except sqlite3.Error as e:
        logging.error(f"Error deleting exam record {exam_id}: {e}", exc_info=True)
        return False

# --- PDF Processing Functions ---

def load_and_split_pdf(pdf_path: str,
                       chunk_size: int = DEFAULT_CHUNK_SIZE,
                       chunk_overlap: int = DEFAULT_CHUNK_OVERLAP) -> List[Document]:
    """Loads a PDF, splits it into chunks, and returns a list of LangChain Documents."""
    try:
        loader = PyPDFLoader(pdf_path)
        pages = loader.load() # Loads pages as individual documents

        if not pages:
             logging.warning(f"No pages loaded from PDF: {pdf_path}")
             return []

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            add_start_index=True, # Helps in tracing back the source
        )

        chunks = text_splitter.split_documents(pages)

        # Add metadata to chunks for source tracking
        pdf_basename = os.path.basename(pdf_path)
        for i, chunk in enumerate(chunks):
            original_page_num = chunk.metadata.get('page', None) # Default to None
            page_info = f' - Sayfa {original_page_num + 1}' if isinstance(original_page_num, int) else ''

            chunk.metadata["source_pdf"] = pdf_basename
            chunk.metadata["chunk_index"] = i
            # Keep the original page number (0-indexed) if available
            if isinstance(original_page_num, int):
                 chunk.metadata["source_page_index"] = original_page_num
            # Create a readable source string for context
            chunk.metadata["source"] = f"{pdf_basename}{page_info}, Parça {i}"

        logging.info(f"Successfully loaded and split PDF '{pdf_basename}' into {len(chunks)} chunks.")
        return chunks

    except Exception as e:
        logging.error(f"Error loading/splitting PDF '{os.path.basename(pdf_path)}': {e}", exc_info=True)
        return [] # Return empty list on error

# --- Utility for saving uploaded file ---
def save_uploaded_file(uploaded_file) -> Optional[str]:
    """Saves the uploaded Streamlit file to the designated folder.
       Returns the full path to the saved file, or None on error.
    """
    if not uploaded_file:
        logging.warning("save_uploaded_file called with no file.")
        return None
    try:
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        # Sanitize filename (optional but recommended)
        # filename = secure_filename(uploaded_file.name) # Requires werkzeug
        filename = os.path.basename(uploaded_file.name) # Basic sanitization
        save_path = os.path.join(UPLOAD_FOLDER, filename)

        with open(save_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        logging.info(f"File '{filename}' saved successfully to '{save_path}'")
        return save_path
    except Exception as e:
        logging.error(f"Error saving uploaded file '{uploaded_file.name}': {e}", exc_info=True)
        return None

# --- Exam Export Functions ---

def _parse_exam_data(exam_data_str: str) -> Optional[Dict]:
    """Parses the JSON string (potentially with markdown fences) into a dictionary."""
    if not exam_data_str:
        return None
    try:
        # Clean potential markdown code fences
        cleaned_response = exam_data_str.strip()
        if cleaned_response.startswith("```json"):
            cleaned_response = cleaned_response[len("```json"):].strip()
        if cleaned_response.endswith("```"):
            cleaned_response = cleaned_response[:-len("```")]

        parsed_data = json.loads(cleaned_response)
        # Basic validation
        if isinstance(parsed_data, dict) and "sorular" in parsed_data and "cevap_anahtari" in parsed_data:
             return parsed_data
        else:
             logging.warning(f"Parsed JSON data lacks expected structure: {cleaned_response[:100]}...")
             # Return structure that export functions can handle
             return {"sorular": parsed_data.get("sorular", []), "cevap_anahtari": parsed_data.get("cevap_anahtari", {})}

    except json.JSONDecodeError as e:
        logging.error(f"Error decoding JSON for export: {e}. Data: {exam_data_str[:100]}...", exc_info=True)
        return None
    except Exception as e:
        logging.error(f"Unexpected error parsing exam data: {e}", exc_info=True)
        return None

def create_pdf_exam(exam_data_str: str, topic: str, pdf_name: Optional[str]) -> Optional[io.BytesIO]:
    """Creates a PDF document for the exam in memory. Returns None on error."""
    exam_data = _parse_exam_data(exam_data_str)
    if not exam_data:
        logging.error("Cannot create PDF exam due to parsing errors.")
        return None

    buffer = io.BytesIO()
    try:
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                                leftMargin=0.75*inch, rightMargin=0.75*inch,
                                topMargin=0.75*inch, bottomMargin=0.75*inch)
        styles = getSampleStyleSheet()
        story = []

        # Define styles
        title_style = styles['h1']
        title_style.alignment = TA_CENTER
        sub_title_style = styles['h3']
        sub_title_style.alignment = TA_CENTER
        question_style = styles['Normal']
        option_style = styles['Normal']
        option_style.leftIndent = 0.5*inch
        answer_style = styles['Normal']
        answer_style.alignment = TA_LEFT

        # Build PDF content
        story.append(Paragraph(f"{topic} Sınavı", title_style))
        story.append(Spacer(1, 0.2*inch))
        if pdf_name:
            story.append(Paragraph(f"(Kaynak Belge: {pdf_name})", sub_title_style))
            story.append(Spacer(1, 0.3*inch))
        else:
             story.append(Spacer(1, 0.1*inch))

        # Add Questions
        questions = exam_data.get("sorular", [])
        for i, q in enumerate(questions):
            q_num = q.get('soru_no', i + 1)
            q_text = q.get('soru_metni', 'Soru metni bulunamadı').replace('\n', '<br/>')
            story.append(Paragraph(f"<b>{q_num}.</b> {q_text}", question_style))
            story.append(Spacer(1, 0.1*inch))

            options = q.get("secenekler", {})
            for key, value in sorted(options.items()): # Sort options for consistency
                option_text = value.replace('\n', '<br/>')
                story.append(Paragraph(f"<b>{key})</b> {option_text}", option_style))
            story.append(Spacer(1, 0.2*inch))

        # Add page break before answers if questions exist
        if questions:
            story.append(PageBreak())

        # Add Answer Key
        answers = exam_data.get("cevap_anahtari", {})
        if answers:
             story.append(Paragraph("Cevap Anahtarı", title_style))
             story.append(Spacer(1, 0.2*inch))

             answer_lines = [f"<b>{num}:</b> {ans}" for num, ans in sorted(answers.items())]
             for line in answer_lines:
                 story.append(Paragraph(line, answer_style))
                 story.append(Spacer(1, 0.1*inch))

        doc.build(story)
        buffer.seek(0)
        logging.info(f"Successfully created PDF exam for topic: {topic}")
        return buffer

    except Exception as e:
        logging.error(f"Error building PDF exam for topic '{topic}': {e}", exc_info=True)
        return None

def create_docx_exam(exam_data_str: str, topic: str, pdf_name: Optional[str]) -> Optional[io.BytesIO]:
    """Creates a DOCX document for the exam in memory. Returns None on error."""
    exam_data = _parse_exam_data(exam_data_str)
    if not exam_data:
        logging.error("Cannot create DOCX exam due to parsing errors.")
        return None

    buffer = io.BytesIO()
    try:
        document = DocxDocument()

        # Build DOCX content
        document.add_heading(f"{topic} Sınavı", level=1)
        if pdf_name:
            document.add_heading(f"(Kaynak Belge: {pdf_name})", level=3)
        document.add_paragraph() # Spacer

        # Add Questions
        questions = exam_data.get("sorular", [])
        for i, q in enumerate(questions):
            q_num = q.get('soru_no', i + 1)
            q_text = q.get('soru_metni', 'Soru metni bulunamadı')
            p = document.add_paragraph()
            p.add_run(f"{q_num}. ").bold = True
            p.add_run(q_text)

            options = q.get("secenekler", {})
            for key, value in sorted(options.items()):
                p_opt = document.add_paragraph(style='List Paragraph')
                p_opt.paragraph_format.left_indent = Inches(0.5)
                p_opt.add_run(f"{key}) ").bold = True
                p_opt.add_run(value)

            document.add_paragraph() # Spacer after each question

        # Add page break before answers if questions exist
        if questions:
            document.add_page_break()

        # Add Answer Key
        answers = exam_data.get("cevap_anahtari", {})
        if answers:
             document.add_heading("Cevap Anahtarı", level=1)
             document.add_paragraph()

             for num, ans in sorted(answers.items()):
                 p_ans = document.add_paragraph()
                 p_ans.add_run(f"{num}: ").bold = True
                 p_ans.add_run(ans)

        # Save to buffer
        document.save(buffer)
        buffer.seek(0)
        logging.info(f"Successfully created DOCX exam for topic: {topic}")
        return buffer

    except Exception as e:
        logging.error(f"Error building DOCX exam for topic '{topic}': {e}", exc_info=True)
        return None 